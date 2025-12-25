import pandas as pd
from docx import Document
from openai import OpenAI, APIError
import json
import io
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import re

# ==========================================
# 1. 核心工具类 (Backend Logic)
# ==========================================

class FileParser:
    """处理文件读取与导出"""
    @staticmethod
    def extract_text(uploaded_file):
        filename = uploaded_file.name.lower()
        content = ""
        try:
            if filename.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(uploaded_file)
                content = df.to_string(index=False)
            elif filename.endswith(('.docx', '.doc')):
                doc = Document(uploaded_file)
                content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            else: # txt
                content = uploaded_file.getvalue().decode("utf-8")
        except Exception as e:
            return f"Error reading file: {str(e)}"
        return content

    @staticmethod
    def optimize_text(text):
        nan_string_to_remove = "NaN" 
        text = re.sub(r'(NaN\s*)+', '', text)
        text = re.sub(r'[a-zA-Z]', '', text)
        return text

    @staticmethod
    def generate_word(text):
        doc = Document()
        for line in text.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio

class SimpleRAG:
    """简易 RAG 引擎"""
    def __init__(self, client, embedding_model="text-embedding-3-small"):
        self.client = client
        self.embedding_model = embedding_model
        self.chunks = []
        self.embeddings = []
        
    def ingest(self, files):
        self.chunks = []
        raw_text = ""
        for f in files:
            raw_text += FileParser.extract_text(f) + "\n"
        step = 2000
        self.chunks = [raw_text[i:i+step] for i in range(0, len(raw_text), step) if raw_text[i:i+step].strip()]
        if self.chunks:
            response = self.client.embeddings.create(input=self.chunks, model=self.embedding_model)
            self.embeddings = [item.embedding for item in response.data]

    def retrieve(self, query, top_k=3):
        if not self.embeddings:
            return "（未上传知识库，无背景信息）"
        q_resp = self.client.embeddings.create(input=[query], model=self.embedding_model)
        q_vec = q_resp.data[0].embedding
        sim_matrix = cosine_similarity([q_vec], self.embeddings)
        top_indices = np.argsort(sim_matrix[0])[::-1][:top_k]
        results = [self.chunks[i] for i in top_indices]
        return "\n---\n".join(results)

class AIAgent:
    """处理所有 LLM 交互"""
    def __init__(self, client, model_name="gpt-4o"):
        self.client = client
        self.model = model_name

    def run_translation(self, text, context, prompt, api_log=None):
        sys_prompt = f"""
        你是一个翻译引擎。
        【背景知识】: {context}
        【用户指令】: {prompt}
        请生成两个版本的翻译：
        1. version_precise: 忠实原文，术语精准，直译为主。
        2. version_fluent: 本土化表达，营销口吻，流畅自然。
        必须返回严格的 JSON 格式: {{"v1": "...", "v2": "..."}}
        """
        if api_log is not None:
            api_log.append(f"--- Translation Request ---\n{text[:200]}...")
        try:
            res = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": f"待翻译文本:\n{text}\n\n{sys_prompt}"}],
                response_format={"type": "json_object"}
            )
            response_content = res.choices[0].message.content
            if api_log is not None:
                api_log.append(f"--- Translation Response ---\n{response_content}")
            return json.loads(response_content)
        except Exception as e:
            return {"error": str(e)}

    def run_review(self, text, context, trans_data, api_log=None):
        judge_prompt = f"""
        【背景】: {context}
        【原文】: {text}
        【版本1】: {trans_data.get('v1')}
        【版本2】: {trans_data.get('v2')}
        请对比两者，选出更好的版本。
        必须返回 JSON: {{"best_version": "v1" 或 "v2", "reason": "...", "suggestion": "..."}}
        """
        try:
            res = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": judge_prompt}],
                response_format={"type": "json_object"}
            )
            return json.loads(res.choices[0].message.content)
        except Exception as e:
            return {"best_version": "v1", "reason": f"Error: {str(e)}", "suggestion": ""}

    def run_qa_check(self, text, rules):
        sys_prompt = f"""
        QA质检标准: {rules}
        如果通过返回 {{"status": "PASS"}}，否则返回 {{"status": "FAIL", "reason": "...", "fix_suggestion": "..."}}
        """
        try:
            res = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": f"待检文本: {text}\n\n{sys_prompt}"}],
                response_format={"type": "json_object"}
            )
            return json.loads(res.choices[0].message.content)
        except Exception as e:
            return {"status": "FAIL", "reason": str(e), "fix_suggestion": ""}
